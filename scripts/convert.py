#!/usr/bin/env python3
"""
convert.py  --xlsx <path> --template <path> --output <path>
Reads 'All' tab from the SLA spreadsheet and fills a copy of the LogDTE docx template.
"""
import sys
import argparse
import copy
import json
from datetime import datetime
from collections import Counter

import openpyxl
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
parser = argparse.ArgumentParser()
parser.add_argument("--xlsx", required=True)
parser.add_argument("--template", required=True)
parser.add_argument("--output", required=True)
args = parser.parse_args()

# ---------------------------------------------------------------------------
# 1. Read XLSX "All" tab
# ---------------------------------------------------------------------------
wb = openpyxl.load_workbook(args.xlsx, data_only=True)
ws = wb["All"]

# --- Summary section (rows 3-10, col D=idx 3, col E=idx 4) ---
summary_map = {}
for row in ws.iter_rows(min_row=3, max_row=10, values_only=True):
    label = row[3]
    value = row[4]
    if label:
        summary_map[label] = value if value is not None else 0

total_tiket       = int(summary_map.get("Total Tiket", 0))
total_pending     = int(summary_map.get("Total Pending", 0))
total_selesai     = int(summary_map.get("Total Tiket Yang Selesai", 0))
total_met_resp    = int(summary_map.get("Total Met Response", 0))
total_met_resol   = int(summary_map.get("Total Met Resolution", 0))
pct_met_resp_raw  = summary_map.get("% Met Response", 0)
pct_met_resol_raw = summary_map.get("% Met Resolution", 0)

def fmt_pct(v):
    try:
        f = float(v)
        return f"{f*100:.0f}%" if f <= 1 else f"{f:.0f}%"
    except Exception:
        return str(v)

pct_met_resp  = fmt_pct(pct_met_resp_raw)
pct_met_resol = fmt_pct(pct_met_resol_raw)

# --- Detail section (from row 16, where col C = row number) ---
# Col indices: 2=No, 3=NoTiket, 4=Ringkasan, 5=Rincian, 6=Pemohon, 7=Penyebab,
#              8=Solusi, 9=Tipe, 10=TanggalLaporan, 11=PIC, 12=Vendor,
#              13=Status, 14=WaktuDeteksi, 15=WaktuRespon, 16=WaktuResolusi,
#              17=SLARespon, 18=SLAResolusi, 19=SLAResponNum, 20=SLAResolNum, 21=Pending
tickets = []
for row in ws.iter_rows(min_row=16, max_row=ws.max_row, values_only=True):
    no = row[2]
    if no is None or not isinstance(no, (int, float)):
        continue
    tgl = row[10]
    tgl_str = tgl.strftime("%d/%m/%Y") if isinstance(tgl, datetime) else (str(tgl) if tgl else "")
    tickets.append({
        "no":        int(no),
        "tiket":     str(row[3] or ""),
        "ringkasan": str(row[4] or ""),
        "rincian":   str(row[5] or ""),
        "pemohon":   str(row[6] or ""),
        "penyebab":  str(row[7] or ""),
        "solusi":    str(row[8] or ""),
        "tipe":      str(row[9] or ""),
        "tanggal":   tgl_str,
        "pic":       str(row[11] or ""),
        "vendor":    str(row[12] or ""),
        "status":    str(row[13] or ""),
        "sla_respon":  str(row[17] or ""),
        "sla_resol":   str(row[18] or ""),
    })

# Sort by PIC (worker id), then by ticket number
tickets.sort(key=lambda t: (t["pic"].lower(), t["no"]))

# Re-number after sort
for i, t in enumerate(tickets, 1):
    t["no"] = i

# --- Top 5 Ringkasan ---
ringkasan_counter = Counter(t["ringkasan"] for t in tickets if t["ringkasan"])
top5 = ringkasan_counter.most_common(5)

# --- Derive Periode and Tanggal from tickets ---
# Periode = month+year of earliest ticket date
# Tanggal = today (report generation date)
MONTHS_ID = {
    1: "Januari", 2: "Februari", 3: "Maret", 4: "April",
    5: "Mei", 6: "Juni", 7: "Juli", 8: "Agustus",
    9: "September", 10: "Oktober", 11: "November", 12: "Desember"
}
dates = []
for row in ws.iter_rows(min_row=16, max_row=ws.max_row, values_only=True):
    no = row[2]
    if no is None or not isinstance(no, (int, float)):
        continue
    tgl = row[10]
    if isinstance(tgl, datetime):
        dates.append(tgl)

if dates:
    earliest = min(dates)
    periode = f"{MONTHS_ID[earliest.month]} {earliest.year}"
else:
    periode = "Agustus 2025"

# Count incident vs service request
tipe_counter = Counter(t["tipe"] for t in tickets)
incident_count = sum(v for k, v in tipe_counter.items() if "incident" in k.lower() or k.upper() == "IN")
sr_count = sum(v for k, v in tipe_counter.items()
               if "work order" in k.lower() or "service request" in k.lower()
               or "change request" in k.lower() or k.upper() == "SR")
# fallback
if incident_count == 0 and sr_count == 0:
    incident_count = sum(v for k, v in tipe_counter.items() if "INC" in k.upper())
    sr_count = total_tiket - incident_count

# ---------------------------------------------------------------------------
# Helper: set cell text preserving runs/formatting  (first paragraph only)
# ---------------------------------------------------------------------------
def set_cell_text(cell, text):
    """Replace all paragraphs in cell with a single paragraph containing text."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = str(text)
            return
        # no runs – add one
        run = p.add_run(str(text))
        return
    # no paragraphs
    p = cell.add_paragraph(str(text))


def get_first_run_or_create(para):
    if para.runs:
        return para.runs[0]
    return para.add_run()


# ---------------------------------------------------------------------------
# 2. Fill the docx template
# ---------------------------------------------------------------------------
doc = Document(args.template)

# --- 2a. Update header: Periode and Tanggal ---
today = datetime.now()
tanggal_str = f"{today.day} {MONTHS_ID[today.month]} {today.year}"

for section in doc.sections:
    ht = section.header.tables[0]
    # Row 1, Col 6 = "Agustus 2025" -> periode
    set_cell_text(ht.rows[1].cells[6], periode)
    # Row 2, Col 6 = "15 September 2025" -> tanggal
    set_cell_text(ht.rows[2].cells[6], tanggal_str)

# --- 2b. Fill detail table (Table 0) ---
# Table has header rows [0] and [1], then data rows starting at [2]
# The template has 126 pre-numbered empty rows (rows 3 to 128)
t0 = doc.tables[0]
DATA_START_ROW = 2  # 0-indexed

def fill_detail_row(table_row, ticket):
    cols = table_row.cells
    # Col 0: No
    set_cell_text(cols[0], ticket["no"])
    # Col 1: No Tiket
    set_cell_text(cols[1], ticket["tiket"])
    # Col 2: Ringkasan
    set_cell_text(cols[2], ticket["ringkasan"])
    # Col 3: Rincian
    set_cell_text(cols[3], ticket["rincian"])
    # Col 4: Pemohon
    set_cell_text(cols[4], ticket["pemohon"])
    # Col 5: Solusi
    set_cell_text(cols[5], ticket["solusi"])
    # Col 6: Tipe
    set_cell_text(cols[6], ticket["tipe"])
    # Col 7: Tanggal
    set_cell_text(cols[7], ticket["tanggal"])
    # Col 8: PIC
    set_cell_text(cols[8], ticket["pic"])
    # Col 9: SLA Respon
    set_cell_text(cols[9], ticket["sla_respon"])
    # Col 10: SLA Resolusi
    set_cell_text(cols[10], ticket["sla_resol"])

# Clone the template data row to use as prototype
template_data_row = t0.rows[DATA_START_ROW]

num_existing = len(t0.rows) - DATA_START_ROW  # rows available for data
num_tickets = len(tickets)

if num_tickets <= num_existing:
    # Fill existing rows then clear the rest
    for i, ticket in enumerate(tickets):
        fill_detail_row(t0.rows[DATA_START_ROW + i], ticket)
    # Clear unused rows
    for i in range(num_tickets, num_existing):
        row = t0.rows[DATA_START_ROW + i]
        for cell in row.cells:
            set_cell_text(cell, "")
else:
    # Fill all existing rows, then add more rows by cloning
    for i in range(num_existing):
        if i < num_tickets:
            fill_detail_row(t0.rows[DATA_START_ROW + i], tickets[i])
    # Add extra rows
    for i in range(num_existing, num_tickets):
        # Clone the last row XML
        new_row = copy.deepcopy(template_data_row._tr)
        t0._tbl.append(new_row)
        fill_detail_row(t0.rows[DATA_START_ROW + i], tickets[i])

# --- 2c. Fill summary table (Table 1) ---
t1 = doc.tables[1]

def set_t1(row_idx, col_idx, text):
    set_cell_text(t1.rows[row_idx].cells[col_idx], text)

set_t1(0, 3, str(total_tiket))         # Total tiket
set_t1(1, 3, str(total_pending))       # Tiket pending
set_t1(2, 3, str(total_selesai))       # Total Tiket Yang Selesai
set_t1(3, 3, str(incident_count))      # Incident
set_t1(4, 3, str(sr_count))            # Service Request
set_t1(5, 3, str(total_met_resp))      # SLA respon yang MET
set_t1(6, 3, str(total_met_resol))     # SLA resolusi yang MET
set_t1(7, 3, pct_met_resp)             # % pencapaian SLA respon
set_t1(8, 3, pct_met_resol)            # % pencapaian SLA resolusi

# Top 5 Ringkasan (rows 1-5, cols 5=rank, 6=ringkasan, 7=colon, 8=count)
for rank, (ring, cnt) in enumerate(top5, 1):
    row_idx = rank  # rows 1-5
    if row_idx < len(t1.rows):
        set_t1(row_idx, 5, str(rank))
        set_t1(row_idx, 6, ring)
        set_t1(row_idx, 7, ":")
        set_t1(row_idx, 8, str(cnt))

# Clear remaining top5 slots if fewer than 5
for rank in range(len(top5) + 1, 6):
    row_idx = rank
    if row_idx < len(t1.rows):
        set_t1(row_idx, 5, "")
        set_t1(row_idx, 6, "")
        set_t1(row_idx, 7, "")
        set_t1(row_idx, 8, "")

# ---------------------------------------------------------------------------
# 3. Save
# ---------------------------------------------------------------------------
doc.save(args.output)
print(json.dumps({"success": True, "tickets": num_tickets, "periode": periode}))
